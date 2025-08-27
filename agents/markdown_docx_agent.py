from docx import Document
from docx.shared import Pt
import os
import re

DOCX_FOLDER = "output"
os.makedirs(DOCX_FOLDER, exist_ok=True)
def save_markdown_docx(markdown_text: str, filename: str) -> str:
    doc = Document()

    # Define simple styles
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    lines = markdown_text.strip().splitlines()
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        # **Bold Heading** (exact match)
        match_heading = re.match(r"^\*\*(.+?)\*\*$", line)
        if match_heading:
            heading = match_heading.group(1).strip()
            doc.add_paragraph(heading, style="Heading 2")
            continue

        # Numbered list (remove digit prefix)
        match_numbered = re.match(r"^(\d+)\.\s+(.*)$", line)
        if match_numbered:
            content = match_numbered.group(2).strip()
            doc.add_paragraph(content, style="List Number")
            continue

        # Bullet list (remove bullet prefix)
        if re.match(r"^[-*+]\s+", line):
            content = re.sub(r"^[-*+]\s+", "", line)
            doc.add_paragraph(content, style="List Bullet")
            continue

        # Otherwise plain paragraph
        doc.add_paragraph(line)

    # Save DOCX
    filename = filename if filename.endswith(".docx") else f"{filename}.docx"
    save_path = os.path.join(DOCX_FOLDER, filename)
    doc.save(save_path)
    return save_path
