from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import os
import re
import uuid
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCX_FOLDER = "output"

def apply_custom_style(doc):
    styles = doc.styles
    if "CustomHeading" not in styles:
        heading_style = styles.add_style("CustomHeading", WD_STYLE_TYPE.PARAGRAPH)
        heading_font = heading_style.font
        heading_font.size = Pt(14)
        heading_font.bold = True

def add_markdown_text(doc, text):
    for line in text.split('\n'):
        clean = line.strip()
        if not clean:
            continue

        # Full bold line as heading
        if clean.startswith("**") and clean.endswith("**") and clean.count("**") == 2:
            doc.add_paragraph(clean.strip("* ").strip(), style="CustomHeading")
        else:
            para = doc.add_paragraph()
            parts = re.split(r'(\*\*[^*]+\*\*)', clean)
            for part in parts:
                run = para.add_run(part.strip("*") if part.startswith("**") else part)
                if part.startswith("**") and part.endswith("**"):
                    run.bold = True

def save_to_docx(result_dict, filename):
    doc = Document()
    apply_custom_style(doc)

    # Title
    doc.add_paragraph("Lecture Notes Summary", style="CustomHeading")
    doc.add_paragraph(f"Detected Language: {result_dict['language']}")

    if result_dict.get("converted_text"):
        doc.add_paragraph("Roman Urdu / Converted Text:", style="CustomHeading")
        add_markdown_text(doc, result_dict["converted_text"])

    doc.add_paragraph("Notes in Local Language:", style="CustomHeading")
    add_markdown_text(doc, result_dict["notes_local"])

    doc.add_paragraph("Notes in English:", style="CustomHeading")
    add_markdown_text(doc, result_dict["notes_english"])

    if result_dict.get("summary"):
        doc.add_paragraph("Brief English Summary:", style="CustomHeading")
        add_markdown_text(doc, result_dict["summary"])

    # Save DOCX file
    os.makedirs(DOCX_FOLDER, exist_ok=True)
    filename = filename.replace(".mp3", ".docx").replace(".wav", ".docx")
    save_path = os.path.join(DOCX_FOLDER, filename)
    doc.save(save_path)
    return save_path
    