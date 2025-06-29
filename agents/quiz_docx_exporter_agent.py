import os
import uuid
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


DOCX_FOLDER = 'output'
def apply_custom_style(doc):
    style = doc.styles.add_style('CustomHeading', 1)
    style.font.name = 'Arial'
    style.font.size = Pt(14)
    style.font.bold = True
    style.paragraph_format.space_after = Pt(12)


def add_markdown_text(doc, markdown_text):
    for line in markdown_text.split('\n'):
        line = line.strip()
        if line.startswith("###"):
            doc.add_paragraph(line.replace("###", "").strip(), style="CustomHeading")
        elif line.startswith("**Answer:**"):
            p = doc.add_paragraph()
            p.add_run("Answer: ").bold = True
            p.add_run(line.replace("**Answer:**", "").strip())
        else:
            doc.add_paragraph(line)


def save_quiz_to_docx(markdown_text, filename=None):
    doc = Document()
    apply_custom_style(doc)

    doc.add_paragraph("Generated Quiz", style="CustomHeading")

    add_markdown_text(doc, markdown_text)

    os.makedirs(DOCX_FOLDER, exist_ok=True)

    if not filename:
        filename = f"quiz_{uuid.uuid4().hex[:8]}.docx"
    filename = filename+".docx"
    save_path = os.path.join(DOCX_FOLDER, filename)
    doc.save(save_path)
    return filename
